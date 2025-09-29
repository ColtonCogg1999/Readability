# streamlit_app.py 
# ----------------------------------------------------
# Plain-language Readability & Glossary Checker (EN/ES/FR/PT)
# Backend analysis + Streamlit UI (single file)
# ----------------------------------------------------

import io
import re
import pathlib
import pandas as pd
import streamlit as st
import spacy
import pyphen
from docx import Document

# =============================
# Streamlit / Page Setup
# =============================

st.markdown("<div style='text-align:center;'>", unsafe_allow_html=True)

# Use the screenshot stored in the repo (replace with your actual filename if renamed)
img_path = pathlib.Path(__file__).parent / "Screenshot 2025-09-16 113559.png"
st.image(str(img_path), width=600)

st.markdown("</div>", unsafe_allow_html=True)

st.title("Readability Scorer (EN/ES/FR/PT)")
st.caption(
    "This is Precise Language Solutions' multilingual readability scorer. "
    "We normalize syllable counts for Spanish/French/Portuguese to the English Fry grid using high-end SPW baselines. "
    "This tool also detects passive voice, although certain edge-cases may slip through the cracks.\n\n"
    "Glossary flags are optional (CSV with 'term' and 'preferred').\n\n"
    "Note: All languages are graded on the English Fry chart after normalization."
)
st.markdown("---")

# =============================
# Load spaCy models (cached)
# =============================
@st.cache_resource(show_spinner=False)
def load_models():
    """Load spaCy pipelines for EN, ES, FR, PT with graceful download if missing."""
    def _load(model_name):
        try:
            return spacy.load(model_name)
        except Exception:
            try:
                from spacy.cli import download
                download(model_name)
                return spacy.load(model_name)
            except Exception:
                return None

    nlp_en = _load("en_core_web_sm")    # includes parser (dep-based passive cues)
    nlp_es = _load("es_core_news_sm")
    nlp_fr = _load("fr_core_news_sm")
    nlp_pt = _load("pt_core_news_sm")

    # Fallback: create blank models if downloads fail (sentence splitter with sentencizer)
    if nlp_en is None:
        nlp_en = spacy.blank("en")
        if "sentencizer" not in nlp_en.pipe_names:
            nlp_en.add_pipe("sentencizer")
    if nlp_es is None:
        nlp_es = spacy.blank("es")
        if "sentencizer" not in nlp_es.pipe_names:
            nlp_es.add_pipe("sentencizer")
    if nlp_fr is None:
        nlp_fr = spacy.blank("fr")
        if "sentencizer" not in nlp_fr.pipe_names:
            nlp_fr.add_pipe("sentencizer")
    if nlp_pt is None:
        nlp_pt = spacy.blank("pt")
        if "sentencizer" not in nlp_pt.pipe_names:
            nlp_pt.add_pipe("sentencizer")

    return nlp_en, nlp_es, nlp_fr, nlp_pt

nlp_en, nlp_es, nlp_fr, nlp_pt = load_models()

# =============================
# Minimal normalization for UI input
# =============================
BULLETS = r'\-\u2013\u2014\u2022·•'
BULLET_LINE = rf'(?m)^[\s]*[{BULLETS}]\s*'

def normalize_for_analysis(t: str) -> str:
    if not t:
        return ""
    t = t.replace("\r", "")
    t = re.sub(BULLET_LINE, "", t)                  # remove bullet glyphs at line starts
    t = re.sub(r"[ \t]+", " ", t)                 # collapse spaces (leave \n alone)
    t = re.sub(r"\n{3,}", "\n\n", t)            # limit huge blank blocks
    return t.strip()

# =============================
# Readability helpers (Fry-like bands)
# =============================
# English Fry retained; all languages will map to EN Fry via normalization.

def estimate_fry_grade_en(avg_sent, avg_syll):
    grades = [
        {"grade": "Grade 1", "sent": (14.0, 25.1), "syll": (108, 124)},
        {"grade": "Grade 2", "sent": (11.0, 14.0), "syll": (124, 132)},
        {"grade": "Grade 3", "sent": (9.0, 11.0), "syll": (132, 140)},
        {"grade": "Grade 4", "sent": (7.7, 9.0), "syll": (140, 146)},
        {"grade": "Grade 5", "sent": (6.3, 7.7), "syll": (146, 152)},
        {"grade": "Grade 6", "sent": (5.6, 6.3), "syll": (152, 156)},
        {"grade": "Grade 7", "sent": (5.0, 5.6), "syll": (156, 160)},
        {"grade": "Grade 8", "sent": (4.4, 5.0), "syll": (160, 164)},
        {"grade": "Grade 9", "sent": (3.9, 4.4), "syll": (164, 168)},
        {"grade": "Grade 10", "sent": (3.5, 3.9), "syll": (168, 172)},
        {"grade": "Grade 11", "sent": (3.1, 3.5), "syll": (172, 176)},
        {"grade": "Grade 12", "sent": (2.7, 3.1), "syll": (176, 182)},
        {"grade": "College", "sent": (2.0, 2.7), "syll": (182, 200)},
    ]
    sent_index = syll_index = None
    for i, g in enumerate(grades):
        if g["sent"][0] <= avg_sent <= g["sent"][1]:
            sent_index = i
        if g["syll"][0] <= avg_syll <= g["syll"][1]:
            syll_index = i
    sent_grade = grades[sent_index]["grade"] if sent_index is not None else "Out of Range"
    syll_grade = grades[syll_index]["grade"] if syll_index is not None else "Out of Range"
    if sent_index == syll_index and sent_index is not None:
        combined = grades[sent_index]["grade"]
    elif sent_index is not None and syll_index is not None:
        approx_index = round((sent_index + syll_index) / 2)
        approx_index = min(approx_index, len(grades) - 1)
        combined = f"Approx: {grades[approx_index]['grade']}"
    else:
        combined = f"Out of Range (sent: {avg_sent:.2f}, syll: {avg_syll:.2f})"
    return {
        "Grade Level (Sentences)": sent_grade,
        "Grade Level (Syllables)": syll_grade,
        "Grade Level (Fry)": combined
    }

# High-end SPW baselines (syllables per word) used for normalization → English Fry grid
SPW_BASELINES = {
    "en": 1.35,  # English baseline
    "es": 1.60,  # Spanish baseline
    "fr": 1.45,  # French baseline
    "pt": 1.55,  # Portuguese baseline
}

def adjust_syllables_for_fry(lang: str, avg_syll_per_100: float) -> float:
    """Return English-equivalent syllables/100 by rescaling with high-end SPW baselines."""
    spw_en = SPW_BASELINES["en"]
    spw_lang = SPW_BASELINES.get(lang, spw_en)
    factor = spw_en / spw_lang if spw_lang else 1.0
    return avg_syll_per_100 * factor

# =============================
# Syllable counters (per language)
# =============================

# English (Pyphen + vowel fallback)
_dic_en = pyphen.Pyphen(lang="en_US")  # or "en_GB"
_EN_WORD_RE = r"\b[a-zA-Z']+\b"

def count_syllables_english(text: str) -> int:
    if not text.strip():
        return 0
    words = re.findall(_EN_WORD_RE, text)
    total = 0
    for w in words:
        try:
            s = len(_dic_en.inserted(w).split("-"))
            if 1 <= s <= 8:
                total += s
            else:
                total += max(1, len(re.findall(r"[aeiouyAEIOUY]+", w)))
        except Exception:
            total += max(1, len(re.findall(r"[aeiouyAEIOUY]+", w)))
    return total

# Spanish (Pyphen + vowel fallback)
_dic_es = pyphen.Pyphen(lang='es')
def count_syllables_spanish(text: str) -> int:
    if not text.strip():
        return 0
    words = re.findall(r'\b[a-zA-ZáéíóúüñÁÉÍÓÚÜÑ]+\b', text)
    total = 0
    for w in words:
        try:
            s = len(_dic_es.inserted(w).split('-'))
            if 1 <= s <= 8:
                total += s
            else:
                total += max(1, len(re.findall(r'[aeiouáéíóúüAEIOUÁÉÍÓÚÜ]+', w)))
        except Exception:
            total += max(1, len(re.findall(r'[aeiouáéíóúüAEIOUÁÉÍÓÚÜ]+', w)))
    return total

# French (Pyphen + vowel fallback)
_dic_fr = pyphen.Pyphen(lang='fr_FR')
_FR_WORD_RE = r'\b[a-zA-ZàâäéèêëïîôöùûüÿçÀÂÄÉÈÊËÏÎÔÖÙÛÜŸÇ\-]+\b'
def count_syllables_french(text: str) -> int:
    if not text.strip():
        return 0
    words = re.findall(_FR_WORD_RE, text)
    total = 0
    for w in words:
        try:
            s = len(_dic_fr.inserted(w).split('-'))
            if 1 <= s <= 8:
                total += s
            else:
                total += max(1, len(re.findall(r'[aeiouyàâäéèêëïîôöùûüÿAEIOUYÀÂÄÉÈÊËÏÎÔÖÙÛÜŸ]+', w)))
        except Exception:
            total += max(1, len(re.findall(r'[aeiouyàâäéèêëïîôöùûüÿAEIOUYÀÂÄÉÈÊËÏÎÔÖÙÛÜŸ]+', w)))
    return total

# Portuguese (choose PT-PT by default; swap to pt_BR if needed)
_dic_pt = pyphen.Pyphen(lang='pt_BR')  # or 'pt_BR'
_PT_WORD_RE = r'\b[a-zA-ZáàâãéêíóôõúçÁÀÂÃÉÊÍÓÔÕÚÇ]+\b'
def count_syllables_portuguese(text: str) -> int:
    if not text.strip():
        return 0
    words = re.findall(_PT_WORD_RE, text)
    total = 0
    for w in words:
        try:
            s = len(_dic_pt.inserted(w).split('-'))
            if 1 <= s <= 8:
                total += s
            else:
                total += max(1, len(re.findall(r'[aeiouáàâãéêíóôõúAEIOUÁÀÂÃÉÊÍÓÔÕÚ]+', w)))
        except Exception:
            total += max(1, len(re.findall(r'[aeiouáàâãéêíóôõúAEIOUÁÀÂÃÉÊÍÓÔÕÚ]+', w)))
    return total

# =============================
# Token & sentence helpers
# =============================

def count_words_with_spacy(text: str, nlp) -> int:
    if not text.strip():
        return 0
    try:
        doc = nlp(text)
        return len([t for t in doc if t.is_alpha])
    except Exception:
        return len(re.findall(r'\b\w+\b', text))


def count_sentences(text: str, lang: str = "en", min_words: int = 4) -> int:
    """
    Count sentences using spaCy's segmentation. For EN/ES, the parser helps; for FR/PT, sentencizer may be used.
    Require >= min_words and at least one VERB or AUX; fallback to raw sents if overly strict.
    """
    if not text.strip():
        return 0

    nlp = {"en": nlp_en, "es": nlp_es, "fr": nlp_fr, "pt": nlp_pt}.get(lang, nlp_en)
    try:
        doc = nlp(text)
        sents = list(doc.sents)

        valid = 0
        for s in sents:
            words = [t for t in s if t.is_alpha]
            if len(words) >= min_words and any(t.pos_ in {"VERB", "AUX"} for t in s):
                valid += 1

        if valid <= 2 and len(sents) >= 5:
            return len(sents)
        return valid

    except Exception:
        rough = re.split(r'[.!?]+', text)
        return len([seg for seg in rough if len(re.findall(r'\b\w+\b', seg)) >= min_words])

# =============================
# Passive voice detection
# =============================

def passive_voice_en(text):
    if not text.strip():
        return []
    try:
        doc = nlp_en(text)
        passive_sentences = []

        for sent in doc.sents:
            if sum(1 for t in sent if t.is_alpha) < 4:
                continue

            has_passive = False

            # --- Path 1: dependency-based cues
            if any(t.dep_ in {"nsubjpass", "auxpass", "nsubj:pass", "aux:pass"} for t in sent):
                has_passive = True

            # --- Path 2: conservative fallback
            if not has_passive:
                tokens = list(sent)

                def looks_adjectival(v):
                    if v.pos_ == "ADJ":
                        return True
                    if v.dep_ in {"amod", "acomp", "attr"}:
                        return True
                    if v.tag_ == "VBN" and any(a.dep_ == "cop" and a.lemma_ == "be" for a in v.children):
                        return True
                    return False

                for i, tok in enumerate(tokens):
                    if tok.lemma_ in {"be", "get"} and tok.pos_ in {"AUX", "VERB"}:
                        j, steps = i + 1, 0
                        while j < len(tokens) and steps < 4:
                            nxt = tokens[j]

                            # exclude progressives like "is going", "be impacting"
                            if nxt.tag_ == "VBG":
                                break

                            is_participle = (
                                (nxt.tag_ == "VBN" and nxt.pos_ == "VERB") or
                                ("VerbForm=Part" in nxt.morph and nxt.pos_ == "VERB")
                            )

                            if is_participle:
                                if "Pass" in nxt.morph.get("Voice"):
                                    has_passive = True
                                else:
                                    if not looks_adjectival(nxt) and not any(c.dep_ == "nsubj" for c in nxt.children):
                                        aux_link = (
                                            tok in list(nxt.ancestors) or
                                            any(c == tok for c in nxt.children if c.dep_ in {"aux", "auxpass"})
                                        )
                                        if aux_link:
                                            has_passive = True
                                break

                            if nxt.pos_ in {"ADV", "PART", "AUX", "PRON", "DET", "ADP"} or nxt.text.lower() == "not":
                                j += 1
                                steps += 1
                                continue
                            else:
                                break
                    if has_passive:
                        break

            # --- Guardrails
            if has_passive:
                vbns = [t for t in sent if t.tag_ == "VBN"]
                if (not any(t.text.lower() == "by" for t in sent)) and vbns and all(
                    (t.pos_ == "ADJ") or (t.dep_ in {"amod", "acomp", "attr"}) for t in vbns
                ):
                    has_passive = False

            if has_passive:
                passive_sentences.append(sent.text.strip())

        return passive_sentences

    except Exception:
        return []


def passive_voice_es(text):
    if not text.strip():
        return []
    try:
        doc = nlp_es(text)
        passive_sentences = []
        for sent in doc.sents:
            if sum(1 for t in sent if t.is_alpha) < 4:
                continue
            has_passive = False
            tokens = list(sent)
            for i, token in enumerate(tokens[:-1]):
                if token.lemma_ == "ser" and i + 1 < len(tokens):
                    nxt = tokens[i + 1]
                    if ("VerbForm=Part" in nxt.morph) and (nxt.pos_ in {"VERB","ADJ"}):
                        has_passive = True
                        break
            if not has_passive:
                for i, token in enumerate(tokens):
                    if token.text.lower() == "se":
                        j, steps = i + 1, 0
                        while j < len(tokens) and steps < 5:
                            nxt = tokens[j]
                            if nxt.pos_ in {"VERB","AUX"}:
                                mood = nxt.morph.get("Mood")
                                person = nxt.morph.get("Person")
                                if not ("Sub" in mood or "Imp" in mood) and any(p == "3" for p in person):
                                    has_passive = True
                                    break
                            if nxt.pos_ in {"PRON","AUX","ADV","PART","DET","ADP"} or nxt.text.lower() == "no":
                                j += 1; steps += 1; continue
                            break
                        if has_passive:
                            break
            if not has_passive:
                enclitics = [v for v in tokens if v.pos_ == "VERB" and v.text.lower().endswith("se")]
                for v in enclitics:
                    vf = v.morph.get("VerbForm")
                    mood = v.morph.get("Mood")
                    if "Imp" in mood or "Sub" in mood:
                        continue
                    if not any(x in {"Inf","Part","Ger"} for x in vf):
                        continue
                    if any(tok.pos_ in {"AUX","VERB"} and "Fin" in tok.morph.get("VerbForm") and any(p == "3" for p in tok.morph.get("Person")) for tok in tokens):
                        has_passive = True
                        break
            if has_passive:
                passive_sentences.append(sent.text.strip())
        return passive_sentences
    except Exception:
        return []


def passive_voice_fr(text: str):
    """Heuristic French passive detector.
    Catches:
      1) être + Participle  (auxiliary + past participle)
      2) se faire / se voir + Infinitive (e.g., "se faire vacciner")
      3) avoir + été + past participle
    Skips short fragments; agent "par" optional.
    """
    if not text.strip():
        return []
    try:
        doc = nlp_fr(text)
        out = []
        for sent in doc.sents:
            if sum(1 for t in sent if t.is_alpha) < 4:
                continue
            tokens = list(sent)
            has_passive = False
            # (1) être + Participle
            for i, tok in enumerate(tokens[:-1]):
                if tok.lemma_ == "être" and tok.pos_ in {"AUX", "VERB"}:
                    j, steps = i + 1, 0
                    while j < len(tokens) and steps < 5:
                        nxt = tokens[j]
                        if ("VerbForm=Part" in nxt.morph and nxt.pos_ in {"VERB", "AUX", "ADJ"}):
                            has_passive = True
                            break
                        if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() in {"ne", "pas", "plus", "jamais", "point"}:
                            j += 1; steps += 1; continue
                        break
                    if has_passive:
                        break
            # (2) se faire / se voir + Infinitive
            if not has_passive:
                for i, tok in enumerate(tokens[:-2]):
                    if tok.text.lower() == "se" and tokens[i+1].lemma_ in {"faire", "voir"} and tokens[i+1].pos_ in {"AUX", "VERB"}:
                        j, steps = i + 2, 0
                        while j < len(tokens) and steps < 5:
                            nxt = tokens[j]
                            if nxt.pos_ in {"VERB", "AUX"} and "VerbForm=Inf" in nxt.morph:
                                has_passive = True
                                break
                            if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"}:
                                j += 1; steps += 1; continue
                            break
                    if has_passive:
                        break
            # (3) avoir + été + past participle
            if not has_passive:
                for i, tok in enumerate(tokens[:-2]):
                    if tok.lemma_ == "avoir" and tok.pos_ in {"AUX", "VERB"}:
                        j, steps = i + 1, 0
                        found_ete = False
                        while j < len(tokens) and steps < 5:
                            nxt = tokens[j]
                            if nxt.lemma_ == "être" and "VerbForm=Part" in nxt.morph:
                                found_ete = True
                                j += 1
                                steps += 1
                                break
                            if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() in {"ne", "pas", "plus", "jamais", "point"}:
                                j += 1; steps += 1; continue
                            break
                        if found_ete:
                            while j < len(tokens) and steps < 8:
                                nxt = tokens[j]
                                if ("VerbForm=Part" in nxt.morph and nxt.pos_ in {"VERB", "AUX", "ADJ"}):
                                    has_passive = True
                                    break
                                if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() in {"ne", "pas", "plus", "jamais", "point"}:
                                    j += 1; steps += 1; continue
                                break
                    if has_passive:
                        break
            if has_passive:
                out.append(sent.text.strip())
        return out
    except Exception:
        return []


def passive_voice_pt(text: str):
    """Heuristic Portuguese passive detector.
    Catches:
      1) ser / estar / ficar + Particípio (past participle)
      2) voz passiva sintética: "se" + 3rd person finite verb (indicative preferred)
      3) ter + sido + past participle
      4) 3rd person singular verb + se (enclitic)
    Skips short fragments; tries to avoid reflexive/imperative noise.
    """
    if not text.strip():
        return []
    try:
        doc = nlp_pt(text)
        out = []
        for sent in doc.sents:
            if sum(1 for t in sent if t.is_alpha) < 4:
                continue
            tokens = list(sent)
            has_passive = False
            # (1) ser/estar/ficar + participle
            for i, tok in enumerate(tokens[:-1]):
                if tok.lemma_ in {"ser", "estar", "ficar"} and tok.pos_ in {"AUX", "VERB"}:
                    j, steps = i + 1, 0
                    while j < len(tokens) and steps < 5:
                        nxt = tokens[j]
                        if ("VerbForm=Part" in nxt.morph and nxt.pos_ in {"VERB", "AUX", "ADJ"}):
                            has_passive = True
                            break
                        if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() == "não":
                            j += 1; steps += 1; continue
                        break
                    if has_passive:
                        break
            # (2) voz passiva sintética: se + 3rd person finite VERB/AUX
            if not has_passive:
                for i, tok in enumerate(tokens):
                    if tok.text.lower() == "se":
                        j, steps = i + 1, 0
                        while j < len(tokens) and steps < 5:
                            nxt = tokens[j]
                            if nxt.pos_ in {"VERB", "AUX"} and "Fin" in nxt.morph.get("VerbForm"):
                                mood = nxt.morph.get("Mood")
                                person = nxt.morph.get("Person")
                                if not ("Sub" in mood or "Imp" in mood) and any(p == "3" for p in person):
                                    has_passive = True
                                    break
                            if nxt.pos_ in {"PRON", "AUX", "ADV", "PART", "DET", "ADP"} or nxt.text.lower() == "não":
                                j += 1; steps += 1; continue
                            break
                        if has_passive:
                            break
            # (3) ter + sido + past participle
            if not has_passive:
                for i, tok in enumerate(tokens[:-2]):
                    if tok.lemma_ == "ter" and tok.pos_ in {"AUX", "VERB"}:
                        j, steps = i + 1, 0
                        found_sido = False
                        while j < len(tokens) and steps < 5:
                            nxt = tokens[j]
                            if nxt.lemma_ == "ser" and "VerbForm=Part" in nxt.morph:
                                found_sido = True
                                j += 1
                                steps += 1
                                break
                            if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() == "não":
                                j += 1; steps += 1; continue
                            break
                        if found_sido:
                            while j < len(tokens) and steps < 8:
                                nxt = tokens[j]
                                if ("VerbForm=Part" in nxt.morph and nxt.pos_ in {"VERB", "AUX", "ADJ"}):
                                    has_passive = True
                                    break
                                if nxt.pos_ in {"ADV", "PART", "AUX", "DET", "ADP", "PRON"} or nxt.text.lower() == "não":
                                    j += 1; steps += 1; continue
                                break
                    if has_passive:
                        break
            # (4) 3rd person singular verb + se (enclitic)
            if not has_passive:
                for i, tok in enumerate(tokens):
                    if tok.pos_ in {"VERB", "AUX"} and "Fin" in tok.morph.get("VerbForm"):
                        mood = tok.morph.get("Mood")
                        person = tok.morph.get("Person")
                        number = tok.morph.get("Number")
                        if (not ("Sub" in mood or "Imp" in mood) and
                            any(p == "3" for p in person) and
                            any(n == "Sing" for n in number)):
                            # Check if token ends with -se
                            if tok.text.lower().endswith("-se"):
                                has_passive = True
                                break
                    if has_passive:
                        break
            if has_passive:
                out.append(sent.text.strip())
        return out
    except Exception:
        return []

# =============================
# Glossary checker (language-agnostic)
# =============================

def glossary_flags_with_suggestions(text: str, glossary_df: pd.DataFrame):
    """Flags when 'term' appears instead of 'preferred'. Exact word or phrase matches."""
    if glossary_df is None or glossary_df.empty or not text.strip():
        return []
    text_lower = text.lower()
    tokens = set(re.findall(r'\b\w+\b', text_lower))
    flags = []
    for _, row in glossary_df.iterrows():
        term = str(row.get('term', '')).lower().strip()
        preferred = str(row.get('preferred', '')).strip()
        if not term:
            continue
        if term in tokens:
            flags.append(f"Used '{term}' instead of '{preferred}'")
        if len(term.split()) > 1 and term in text_lower:
            flags.append(f"Used '{term}' instead of '{preferred}'")
    return flags

# =============================
# Analysis wrappers (all grade on EN Fry after normalization)
# =============================

def analyze_english(text: str, glossary_df: pd.DataFrame):
    if not text.strip():
        return {
            "Error": "No text provided for analysis",
            "Word Count": 0, "Sentence Count": 0, "Syllable Count": 0,
            "Avg Sentences per 100 Words": 0, "Avg Syllables per 100 Words": 0,
            "Grade Level": "N/A", "Passive Voice Count": 0, "Passive Voice Sentences": [],
            "Glossary Flags": {"Count": 0, "Messages": []}
        }
    try:
        words = count_words_with_spacy(text, nlp_en)
        sentences = count_sentences(text, lang="en")
        syllables = count_syllables_english(text)
        if words == 0:
            return {"Error": "No valid words found in text"}
        avg_sent = (sentences / words) * 100
        avg_syll = (syllables / words) * 100
        syll_eq = adjust_syllables_for_fry("en", avg_syll)  # EN-normalized (same scale)
        grade = estimate_fry_grade_en(avg_sent, syll_eq)
        passive_sentences = passive_voice_en(text)
        glossary_flags = glossary_flags_with_suggestions(text, glossary_df)
        return {
            "Word Count": words,
            "Sentence Count": sentences,
            "Syllable Count": syllables,
            "Avg Sentences per 100 Words": round(avg_sent, 2),
            "Avg Syllables per 100 Words": round(avg_syll, 2),
            "Grade Level (Fry, EN-normalized)": grade,
            "Passive Voice Count": len(passive_sentences),
            "Passive Voice Sentences": passive_sentences,
            "Glossary Flags": {"Count": len(glossary_flags), "Messages": glossary_flags}
        }
    except Exception as e:
        return {"Error": f"Analysis failed: {str(e)}"}


def analyze_spanish(text: str, glossary_df: pd.DataFrame):
    if not text.strip():
        return {
            "Error": "No text provided for analysis",
            "Word Count": 0, "Sentence Count": 0, "Syllable Count": 0,
            "Avg Sentences per 100 Words": 0, "Avg Syllables per 100 Words": 0,
            "Grade Level": "N/A", "Passive Voice Count": 0, "Passive Voice Sentences": [],
            "Glossary Flags": {"Count": 0, "Messages": []}
        }
    try:
        words = count_words_with_spacy(text, nlp_es)
        sentences = count_sentences(text, lang="es")
        syllables = count_syllables_spanish(text)
        if words == 0:
            return {"Error": "No valid words found in text"}
        avg_sent = (sentences / words) * 100
        avg_syll = (syllables / words) * 100
        syll_eq = adjust_syllables_for_fry("es", avg_syll)
        grade = estimate_fry_grade_en(avg_sent, syll_eq)
        passive_sentences = passive_voice_es(text)
        glossary_flags = glossary_flags_with_suggestions(text, glossary_df)
        return {
            "Word Count": words,
            "Sentence Count": sentences,
            "Syllable Count": syllables,
            "Avg Sentences per 100 Words": round(avg_sent, 2),
            "Avg Syllables per 100 Words": round(avg_syll, 2),
            "Grade Level (Fry, EN-normalized)": grade,
            "Passive Voice Count": len(passive_sentences),
            "Passive Voice Sentences": passive_sentences,
            "Glossary Flags": {"Count": len(glossary_flags), "Messages": glossary_flags}
        }
    except Exception as e:
        return {"Error": f"Analysis failed: {str(e)}"}


def analyze_french(text: str, glossary_df: pd.DataFrame):
    if not text.strip():
        return {
            "Error": "No text provided for analysis",
            "Word Count": 0, "Sentence Count": 0, "Syllable Count": 0,
            "Avg Sentences per 100 Words": 0, "Avg Syllables per 100 Words": 0,
            "Grade Level": "N/A", "Passive Voice Count": 0, "Passive Voice Sentences": [],
            "Glossary Flags": {"Count": 0, "Messages": []}
        }
    try:
        words = count_words_with_spacy(text, nlp_fr)
        sentences = count_sentences(text, lang="fr")
        syllables = count_syllables_french(text)
        if words == 0:
            return {"Error": "No valid words found in text"}
        avg_sent = (sentences / words) * 100
        avg_syll = (syllables / words) * 100
        syll_eq = adjust_syllables_for_fry("fr", avg_syll)
        grade = estimate_fry_grade_en(avg_sent, syll_eq)
        passive_sentences = passive_voice_fr(text)
        glossary_flags = glossary_flags_with_suggestions(text, glossary_df)
        return {
            "Word Count": words,
            "Sentence Count": sentences,
            "Syllable Count": syllables,
            "Avg Sentences per 100 Words": round(avg_sent, 2),
            "Avg Syllables per 100 Words": round(avg_syll, 2),
            "Grade Level (Fry, EN-normalized)": grade,
            "Passive Voice Count": len(passive_sentences),
            "Passive Voice Sentences": passive_sentences,
            "Glossary Flags": {"Count": len(glossary_flags), "Messages": glossary_flags}
        }
    except Exception as e:
        return {"Error": f"Analysis failed: {str(e)}"}


def analyze_portuguese(text: str, glossary_df: pd.DataFrame):
    if not text.strip():
        return {
            "Error": "No text provided for analysis",
            "Word Count": 0, "Sentence Count": 0, "Syllable Count": 0,
            "Avg Sentences per 100 Words": 0, "Avg Syllables per 100 Words": 0,
            "Grade Level": "N/A", "Passive Voice Count": 0, "Passive Voice Sentences": [],
            "Glossary Flags": {"Count": 0, "Messages": []}
        }
    try:
        words = count_words_with_spacy(text, nlp_pt)
        sentences = count_sentences(text, lang="pt")
        syllables = count_syllables_portuguese(text)
        if words == 0:
            return {"Error": "No valid words found in text"}
        avg_sent = (sentences / words) * 100
        avg_syll = (syllables / words) * 100
        syll_eq = adjust_syllables_for_fry("pt", avg_syll)
        grade = estimate_fry_grade_en(avg_sent, syll_eq)
        passive_sentences = passive_voice_pt(text)
        glossary_flags = glossary_flags_with_suggestions(text, glossary_df)
        return {
            "Word Count": words,
            "Sentence Count": sentences,
            "Syllable Count": syllables,
            "Avg Sentences per 100 Words": round(avg_sent, 2),
            "Avg Syllables per 100 Words": round(avg_syll, 2),
            "Grade Level (Fry, EN-normalized)": grade,
            "Passive Voice Count": len(passive_sentences),
            "Passive Voice Sentences": passive_sentences,
            "Glossary Flags": {"Count": len(glossary_flags), "Messages": glossary_flags}
        }
    except Exception as e:
        return {"Error": f"Analysis failed: {str(e)}"}


def _grade_label_for_metric(grade_obj) -> str:
    if isinstance(grade_obj, dict):
        return grade_obj.get("Grade Level (Fry)", "N/A")
    return str(grade_obj or "N/A")

# =============================
# Sidebar controls
# =============================
with st.sidebar:
    st.header("Settings")
    lang = st.radio("Document language", ["English", "Spanish", "French", "Portuguese"], index=0)
    lang_code = {"English": "en", "Spanish": "es", "French": "fr", "Portuguese": "pt"}[lang]
    input_mode = st.radio("Input method", ["Paste text", "Upload .docx"], index=0)
    glossary_file = st.file_uploader("Glossary CSV (optional)", type=["csv"])
    debug_sentences = st.checkbox("Show sentence boundaries (debug)", value=False)

# =============================
# Layout
# =============================

# Input section (full width)
st.subheader("Input / Preview")
text = ""

if input_mode == "Paste text":
    text = st.text_area("Paste content", "", height=200)
else:
    doc_file = st.file_uploader("Word document (.docx)", type=["docx"])
    if doc_file:
        try:
            d = Document(io.BytesIO(doc_file.read()))
            full_text = []
            # paragraphs
            for para in d.paragraphs:
                if para.text and para.text.strip():
                    full_text.append(para.text.strip())
            # tables
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            full_text.append(cell.text.strip())
            text = "\n".join(full_text)
            if text.strip():
                with st.expander("Preview extracted text"):
                    st.code(text[:2000] + "..." if len(text) > 2000 else text)
            else:
                st.warning("No text extracted from document")
        except Exception as e:
            st.error(f"Docx read error: {e}")

# Analysis section (full width)
st.subheader("Analysis")

if not text.strip():
    st.info("Provide input above to see analysis results.")
else:
    # Keep newlines; minimal cleanup only
    text = normalize_for_analysis(text)

    # Load glossary (optional)
    try:
        gloss_df = pd.read_csv(glossary_file) if glossary_file else pd.DataFrame(columns=["term", "preferred"])
        # enforce columns exactly
        if not set(gloss_df.columns.str.lower()) >= {"term", "preferred"}:
            st.warning("Glossary CSV must contain columns exactly: 'term' and 'preferred'. Using empty glossary.")
            gloss_df = pd.DataFrame(columns=["term", "preferred"])
        else:
            # standardize column names
            gloss_df.columns = [c.lower().strip() for c in gloss_df.columns]
    except Exception as e:
        st.error(f"Error reading glossary file: {e}")
        gloss_df = pd.DataFrame(columns=["term", "preferred"])

    # Run analysis
    if lang_code == "en":
        result = analyze_english(text, gloss_df)
    elif lang_code == "es":
        result = analyze_spanish(text, gloss_df)
    elif lang_code == "fr":
        result = analyze_french(text, gloss_df)
    else:
        result = analyze_portuguese(text, gloss_df)

    if "Error" in result and result.get("Error"):
        st.error(result["Error"])
    else:
        # Top row - 3 metrics
        m1, m2, m3 = st.columns(3)
        m1.metric("Words", f"{result['Word Count']:,}")
        m2.metric("Sentences", f"{result['Sentence Count']:,}")
        m3.metric("Syllables", f"{result['Syllable Count']:,}")

        # Second row - Grade gets full space
        st.metric("Reading Grade Level", _grade_label_for_metric(result.get("Grade Level (Fry, EN-normalized)")))

        # Table of details
        details = {
            "Avg Sentences per 100 Words": result["Avg Sentences per 100 Words"],
            "Avg Syllables per 100 Words": result["Avg Syllables per 100 Words"],
            "Passive Voice (Count)": result["Passive Voice Count"],
            "Glossary Flags (Count)": result["Glossary Flags"]["Count"]
        }
        st.table(pd.DataFrame({"Metric": list(details.keys()), "Value": list(details.values())}))

        # Expanders
        with st.expander("Passive voice sentences"):
            if result["Passive Voice Sentences"]:
                for i, s in enumerate(result["Passive Voice Sentences"], 1):
                    st.markdown(f"**{i}.** {s}")
            else:
                st.write("None detected.")

        with st.expander("Glossary flags"):
            msgs = result["Glossary Flags"]["Messages"]
            if msgs:
                st.dataframe(pd.DataFrame({"Message": msgs}), use_container_width=True)
            else:
                st.write("No issues found.")

        # Optional debug: show actual sentence segmentation from spaCy
        if debug_sentences:
            with st.expander("Debug: Sentence boundaries"):
                nlp_map = {"en": nlp_en, "es": nlp_es, "fr": nlp_fr, "pt": nlp_pt}
                nlp = nlp_map[lang_code]
                for i, s in enumerate(nlp(text).sents, 1):
                    st.write(f"{i}. {s.text}")

        st.subheader("Detailed Results:")
        st.json(result)

st.caption(
    "Tip: Pin versions for reproducibility: spacy==3.7.x, en-core-web-sm==3.7.x, es-core-news-sm==3.7.x, "
    "fr-core-news-sm==3.7.x, pt-core-news-sm==3.7.x, pyphen==0.14.0, python-docx==1.1.2, pandas>=2.0. "
    "Normalization baselines (high-end SPW): EN 1.45, ES 1.70, FR 1.55, PT 1.65. All syllables use Pyphen."
)
